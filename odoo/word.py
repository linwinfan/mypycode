#!/usr/bin/python3
# -*- coding: utf-8 -*-

#读取docx中的文本代码示例
import docx
import re

import odoorpc
import sys
import datetime
import base64
from optparse import OptionParser

def parse_args():
  parser = OptionParser()
  
  parser.add_option("", "--odoo", dest="odoo", help="", action="store", type="string", default='172.21.9.53')
  parser.add_option("", "--port", dest="port", help="", action="store", type="int", default=8069)
  parser.add_option("", "--user", dest="user", help="", action="store", type="string", default='linwinfan@163.com')
  parser.add_option("", "--pwd", dest="pwd", help="", action="store", type="string", default='')
  parser.add_option("", "--db", dest="db", help="", action="store", type="string", default='test')
  parser.add_option("", "--method", dest="method", help="", action="store", type="string", default='')
  parser.add_option("", "--args", dest="args", help="", action="store", type="string", default='')
  parser.add_option("", "--arg1", dest="arg1", help="", action="store", type="string", default='')
  parser.add_option("", "--arg2", dest="arg2", help="", action="store", type="string", default='')
  parser.add_option("", "--arg3", dest="arg3", help="", action="store", type="string", default='')
  (options, args) = parser.parse_args()
  
  return options

def login(options):
    # Prepare the connection to the server
    odoo = odoorpc.ODOO(options.odoo,protocol='jsonrpc', port=options.port)

    # Check available databases
    #print(odoo.db.list())

    # Login
    odoo.login(options.db, options.user, options.pwd)
    return odoo

def importdoc(odoo):
    #获取文档
    file=docx.Document("C:\\Users\\cec-lcg\\Desktop\\饶修美-中山市石岐区南江路70号1幢602房-GZ.docx")
    print("段落数:"+str(len(file.paragraphs))) #输出段落数
    file_word = docx.Document()

    #输出每一段的内容
    doctext=''
    for para in file.paragraphs:
        if(para.text.strip()!=''):
            doctext=doctext+para.text+'\n'
    items={}
    #print(doctext)
    #m = re.match('^H.*?(\d+).*?python$', content, re.S)
    #m=re.search(r'[\S|\s|\n]+\n(\w+)（先生/女士）:[\n|\S|\s]，对位于(\w+)的(\w+)价值作初步评估价。',doctext)
    m=re.search(r'[\S|\s|\n]+\n(\w+)（先生/女士）:[\S|\s|\n|.]+对位于(\w+)的(\w+)价值作初步估价[\S|\s|\n|.]+建筑面积(\d+\.\d+)平方米[\S|\s|\n|.]+在价值时点(\d+年\d+月\d+日)的抵押评估价值为[\S|\s|\n|.]+单价：([\d|,]+)元',doctext)
    if(m):
        print(m.groups())   
        items={'x_name':m.groups()[1],'x_zzjl_qsr':m.groups()[0],'x_md':m.groups()[2],'x_jzmj':float(m.groups()[3]),'x_cp_date': datetime.datetime.strptime(m.groups()[4],'%Y年%m月%d日').strftime('%Y-%m-%d'),'x_cp_dj':m.groups()[5].replace(',','')}
    else:
        print(doctext)
        return
    tables = file.tables #获取文件中的表格集
    table = tables[0]#获取文件中的第一个表格
    #print(dir(table))
    temp=table.cell(1,2).text
    fi=temp.find('构')
    print(fi)
    items.update({'x_wykg_jg':temp[0:fi+1]})
    if(fi<len(temp)):
        items.update({'x_wykg_zcs':temp[fi+1:]})
    items.update({'x_wykg_slss':table.cell(1,5).text})
    items.update({'x_wykg_syzk':table.cell(1,7).text})
    items.update({'x_wykg_hx':table.cell(2,2).text})
    items.update({'x_wykg_hjjg':table.cell(3,5).text})
    items.update({'x_wykg_bxll':table.cell(3,7).text})
    items.update({'x_wykg_ll':table.cell(4,2).text})
    items.update({'x_wykg_tfcg':table.cell(4,5).text})

    items.update({'x_wykg_cx':table.cell(3,2).text})
    items.update({'x_wykg_zxqk':table.cell(5,2).text})

    items.update({'x_zzjl_bh':table.cell(6,2).text})
    items.update({'x_zzjl_qsly':table.cell(6,7).text})
    items.update({'x_zzjl_tdcrjqk':table.cell(7,2).text})
    items.update({'x_ptss_desc':table.cell(8,2).text})
    print(items)
    #图片处理
    shapes = file.inline_shapes #获取文件中的表格集
    document_part = file.part
    i=0
    images=['x_img_loc_s_150_150','x_img_wg','x_img_xqhj','x_img_mph','x_img_slzk_1','x_img_slzk_2','x_img_slzk_3','x_img_slzk_4','','x_img_yt','x_img_ytjg']
    for s in shapes:
        print (s.height.cm,s.width.cm,s._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name)
        blip = s._inline.graphic.graphicData.pic.blipFill.blip
        rId = blip.embed
        image_part = document_part.related_parts[rId]
        #with open('C:\\Users\\cec-lcg\\Desktop\\'+s._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name,'wb') #as fp:
        #    fp.write(image_part._blob) 
        items.update({images[i]:base64.b64encode(image_part._blob).decode('ascii')})
        i=i+1
    

    pg_item=odoo.env['x_pg_item']
    pg_item.create(items)
    """ for i in range(1,len(table.rows)):#从表格第二行开始循环读取表格数据
        for c in range(1,len(table.columns)):
            result ='row:%d col:%d value:%s' %(i,c, table.cell(i,c).text )
            print(result)
        #cell(i,0)表示第(i+1)行第1列数据，以此类推
        #print(result)
    document_part = file.part
    """


    #输出段落编号及段落内容
    """ para_data = []
    for i in range(len(file.paragraphs)):
        # for j in map(lambda x:x.split(' '),file.paragraphs[i].text):
        para_single = file.paragraphs[i].text.split(' ')
        while '' in para_single:  # 移除空格
            para_single.remove('')
        # para_data.append(para_single)
        for data_number in range(len(para_single)):
            data_num = re.findall(r"\d", para_single[data_number])
            data_num = ''.join(data_num)
            para_data.append(data_num + '    ')
    file_word.add_paragraph(para_data)
    file_word.save("E:\\python_word\\number.docx") """
if __name__ == "__main__":
    options=parse_args()
    odoo=login(options)
    importdoc(odoo)